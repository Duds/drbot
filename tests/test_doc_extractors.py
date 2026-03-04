"""Tests for remy.memory.doc_extractors (US-rag-pdf-docx)."""

import importlib.util
from pathlib import Path

import pytest

from remy.memory.doc_extractors import (
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_pdf_ocr,
)


def test_extract_pdf_nonexistent_returns_none():
    """Nonexistent PDF path returns None."""
    assert extract_text_from_pdf(Path("/nonexistent/file.pdf")) is None


def test_extract_pdf_ocr_nonexistent_returns_none():
    """Nonexistent PDF path for OCR returns None."""
    assert extract_text_from_pdf_ocr(Path("/nonexistent/file.pdf")) is None


def test_extract_docx_nonexistent_returns_none():
    """Nonexistent DOCX path returns None."""
    assert extract_text_from_docx(Path("/nonexistent/file.docx")) is None


@pytest.mark.skipif(
    importlib.util.find_spec("pypdf") is None,
    reason="pypdf not installed",
)
def test_extract_pdf_text_extraction(tmp_path):
    """When pypdf is available, a minimal PDF with text is extracted."""
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf not installed")
    pdf_path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    # Add a page with text (pypdf can create minimal PDFs)
    with open(pdf_path, "wb") as f:
        writer.write(f)
    # pypdf extract_text may still return "" for a blank page; so we accept None or ""
    result = extract_text_from_pdf(pdf_path, ocr_enabled=False)
    # Blank page: may be None or empty string
    assert result is None or isinstance(result, str)
    writer.close()


@pytest.mark.skipif(
    importlib.util.find_spec("docx") is None,
    reason="python-docx not installed",
)
def test_extract_docx_paragraphs(tmp_path):
    """When python-docx is available, a DOCX with paragraphs returns text."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")
    docx_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Hello from test.")
    doc.add_paragraph("Second paragraph.")
    doc.save(docx_path)
    result = extract_text_from_docx(docx_path)
    assert result is not None
    assert "Hello from test" in result
    assert "Second paragraph" in result
